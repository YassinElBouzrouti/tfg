from pathlib import Path
import re

from docx import Document
from docx.shared import Pt, Inches


ROOT = Path(r"C:\Users\Yassin\OneDrive\Escritorio\gym_artes_marciales")
SOURCE_MD = ROOT / "MEMORIA_TFG_DAW_YASSINS_GYM.md"
OUTPUT_DOCX = ROOT / "Esqueleto_Memoria_TFG_online_v1.docx"
CAPTURES_DIR = ROOT / "capturas_memoria"


FIGURE_IMAGE_MAP = {
    1: "figure_1_architecture.png",
    2: "figure_2_navigation.png",
    3: "figure_3_network.png",
    4: "figure_4_er.png",
    5: "figure_5_components.png",
    6: "01_home_pc.png",
    7: "05_home_tablet.png",
    8: "06_home_mobile.png",
    9: "02_register_pc.png",
    10: "03_login_pc.png",
    11: "04_admin_access_pc.png",
    12: "07_member_dashboard_full.png",
    13: "07_member_section_6.png",
    14: "07_member_section_3.png",
    15: "07_member_section_2.png",
    16: "08_admin_dashboard_full.png",
    17: "08_admin_dashboard_full.png",
    18: "08_admin_dashboard_full.png",
    19: "09_admin_member_profile.png",
    20: "figure_20_responsive.png",
    21: "figure_21_gantt.png",
}


def clean_inline_md(text: str) -> str:
    cleaned = text.strip()
    cleaned = cleaned.replace("**", "")
    cleaned = cleaned.replace("`", "")
    return cleaned


def is_table_line(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|")


def is_separator_row(line: str) -> bool:
    content = line.strip().strip("|").strip()
    if not content:
        return False
    return all(ch in "-:| " for ch in content)


def parse_table_row(line: str) -> list[str]:
    cells = [clean_inline_md(cell) for cell in line.strip().strip("|").split("|")]
    return [c.strip() for c in cells]


def set_default_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_heading_from_markdown(document: Document, line: str) -> bool:
    if not line.startswith("#"):
        return False
    level = len(line) - len(line.lstrip("#"))
    text = clean_inline_md(line[level:].strip())
    if level > 4:
        level = 4
    document.add_heading(text, level=level)
    return True


def add_figure_if_available(document: Document, line: str) -> bool:
    match = re.match(r"Figura\s+(\d+)\.\s*(.+)", line.strip())
    if not match:
        return False

    fig_number = int(match.group(1))
    image_name = FIGURE_IMAGE_MAP.get(fig_number)

    p = document.add_paragraph()
    run = p.add_run(clean_inline_md(line.strip()))
    run.bold = True

    if image_name:
        image_path = CAPTURES_DIR / image_name
        if image_path.exists():
            document.add_picture(str(image_path), width=Inches(6.3))
    return True


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for idx, h in enumerate(headers):
        run = header_cells[idx].paragraphs[0].add_run(h)
        run.bold = True

    for row in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            if idx < len(row_cells):
                row_cells[idx].text = value


def main() -> None:
    text = SOURCE_MD.read_text(encoding="utf-8")
    lines = text.splitlines()

    document = Document()
    set_default_style(document)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("[INSERTAR CAPTURA:"):
            i += 1
            continue

        if add_heading_from_markdown(document, stripped):
            i += 1
            continue

        if add_figure_if_available(document, stripped):
            i += 1
            continue

        if stripped == "---":
            document.add_page_break()
            i += 1
            continue

        if is_table_line(stripped):
            table_lines = []
            while i < len(lines) and is_table_line(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1

            if len(table_lines) >= 2 and is_separator_row(table_lines[1]):
                headers = parse_table_row(table_lines[0])
                data_rows = [parse_table_row(r) for r in table_lines[2:]]
                add_table(document, headers, data_rows)
            else:
                for raw in table_lines:
                    document.add_paragraph(clean_inline_md(raw))
            continue

        if not stripped:
            document.add_paragraph("")
            i += 1
            continue

        bullet_match = re.match(r"^-\s+(.*)$", stripped)
        numbered_match = re.match(r"^\d+\.\s+(.*)$", stripped)

        if bullet_match:
            document.add_paragraph(clean_inline_md(bullet_match.group(1)), style="List Bullet")
            i += 1
            continue

        if numbered_match:
            document.add_paragraph(clean_inline_md(numbered_match.group(1)), style="List Number")
            i += 1
            continue

        document.add_paragraph(clean_inline_md(line))
        i += 1

    document.save(str(OUTPUT_DOCX))
    print(f"Documento generado: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()

